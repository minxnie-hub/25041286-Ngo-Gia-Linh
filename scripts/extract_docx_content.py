from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def safe_name(value: str) -> str:
    cleaned = re.sub(r"[^\w.-]+", "-", value, flags=re.UNICODE).strip("-")
    return cleaned or "document"


def paragraph_data(paragraph: Paragraph) -> dict:
    text = "".join(paragraph._p.xpath(".//w:t/text()"))
    image_ids = paragraph._p.xpath(".//a:blip/@r:embed")
    images = []
    for relationship_id in image_ids:
        relationship = paragraph.part.rels.get(relationship_id)
        if relationship and relationship.target_ref:
            images.append(Path(relationship.target_ref).name)
    return {
        "type": "paragraph",
        "style": paragraph.style.name if paragraph.style else None,
        "text": text.strip(),
        "images": images,
    }


def table_data(table: Table) -> dict:
    rows = []
    for row in table.rows:
        rows.append([
            "\n".join(
                "".join(paragraph._p.xpath(".//w:t/text()"))
                for paragraph in cell.paragraphs
            ).strip()
            for cell in row.cells
        ])
    return {"type": "table", "rows": rows}


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield paragraph_data(Paragraph(child, document))
        elif child.tag.endswith("}tbl"):
            yield table_data(Table(child, document))


def extract_document(path: Path, output_root: Path) -> dict:
    document = Document(path)
    document_output = output_root / safe_name(path.stem)
    media_output = document_output / "media"
    media_output.mkdir(parents=True, exist_ok=True)

    with ZipFile(path) as archive:
        media_entries = [name for name in archive.namelist() if name.startswith("word/media/")]
        for entry in media_entries:
            target = media_output / Path(entry).name
            with archive.open(entry) as source, target.open("wb") as destination:
                shutil.copyfileobj(source, destination)

    blocks = list(iter_blocks(document))
    payload = {
        "source": str(path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "media_count": len(list(media_output.iterdir())),
        "blocks": blocks,
    }
    (document_output / "content.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return {
        "name": path.name,
        "paragraphs": payload["paragraph_count"],
        "tables": payload["table_count"],
        "media": payload["media_count"],
        "output": str(document_output),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_dir", type=Path)
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    summary = [
        extract_document(path, args.output_dir)
        for path in sorted(args.input_dir.glob("*.docx"))
    ]
    (args.output_dir / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
